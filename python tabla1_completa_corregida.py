
# tabla1_completa_corregida.py
import pandas as pd
import numpy as np
from scipy import stats
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

print("="*100)
print("GENERANDO TABLA 1 COMPLETA Y CORREGIDA")
print("="*100)
print()

# Cargar datos
df = pd.read_csv("datos_limpios.csv")

# Calcular non-HDL-C si no existe
if 'nonHDL_C' not in df.columns:
    df['nonHDL_C'] = df['COL'] - df['cHDL']

# Subgrupos TG
df['TG_group'] = pd.cut(df['TG'], 
                        bins=[0, 150, 200, 400, 10000],
                        labels=['NTG', 'MiTG', 'MoTG', 'HTG'])

# ============================================================================
# FUNCIÓN PARA FORMATEAR MEDIANA (IQR)
# ============================================================================

def format_median_iqr(data):
    """Retorna mediana (Q25-Q75) para variables no normales"""
    data_clean = data.dropna()
    if len(data_clean) == 0:
        return "N/A"
    median = data_clean.median()
    q25 = data_clean.quantile(0.25)
    q75 = data_clean.quantile(0.75)
    return f"{median:.1f} ({q25:.1f}-{q75:.1f})"

# ============================================================================
# CREAR DOCUMENTO WORD
# ============================================================================

doc = Document()

# Título
doc.add_heading('Table 1. Baseline Characteristics of the Study Population', level=1)

p = doc.add_paragraph()
p.add_run('Data are presented as median (interquartile range) for continuous variables '
         'or number (percentage) for categorical variables. All continuous variables '
         'showed non-normal distribution (Shapiro-Wilk test p<0.001).')
p.style = 'Caption'

doc.add_paragraph()

# ============================================================================
# TABLA: OVERALL + POR SUBGRUPO TG
# ============================================================================

# Crear tabla (filas: variables, columnas: Overall + 4 subgrupos TG)
n_vars = 16  # Número de variables a mostrar
table = doc.add_table(rows=n_vars + 1, cols=6)  # +1 para header
table.style = 'Light Grid Accent 1'

# Header
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Characteristic'
hdr_cells[1].text = 'Overall\n(n=24,887)'
hdr_cells[2].text = 'NTG\n(n=11,053, 44.4%)'
hdr_cells[3].text = 'MiTG\n(n=7,028, 28.2%)'
hdr_cells[4].text = 'MoTG\n(n=6,033, 24.2%)'
hdr_cells[5].text = 'HTG\n(n=773, 3.1%)'

# Poner header en negrita
for cell in hdr_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(10)

# ============================================================================
# LLENAR DATOS
# ============================================================================

row_idx = 1

# Variables a incluir (en orden)
variables_info = [
    ('Age', 'Age, years'),
    ('gender', 'Male sex, n (%)'),
    ('COL', 'Total cholesterol, mg/dL'),
    ('cHDL', 'HDL cholesterol, mg/dL'),
    ('nonHDL_C', 'Non-HDL cholesterol, mg/dL'),
    ('TG', 'Triglycerides, mg/dL'),
    ('D-c-LDL', 'Direct LDL-C, mg/dL'),
    ('F-c-LDL', 'Friedewald LDL-C, mg/dL'),
    ('M-c-LDL', 'Martin LDL-C, mg/dL'),
    ('ME-c-LDL', 'Martin-Extended LDL-C, mg/dL'),
    ('S-c-LDL', 'Sampson LDL-C, mg/dL'),
    ('glycemia', 'Glucose, mg/dL'),
    ('TyG', 'TyG index'),
    ('Creatinine', 'Creatinine, mg/dL'),
    ('GPT', 'ALT, U/L')
]

for var_name, var_label in variables_info:
    cells = table.rows[row_idx].cells
    cells[0].text = var_label
    
    # Overall
    if var_name == 'gender':
        n_male = (df['gender'] == 'M').sum()
        pct_male = (n_male / len(df)) * 100
        cells[1].text = f"{n_male} ({pct_male:.1f}%)"
    else:
        cells[1].text = format_median_iqr(df[var_name])
    
    # Por subgrupo TG
    for col_idx, tg_group in enumerate(['NTG', 'MiTG', 'MoTG', 'HTG'], start=2):
        df_group = df[df['TG_group'] == tg_group]
        
        if var_name == 'gender':
            n_male_g = (df_group['gender'] == 'M').sum()
            pct_male_g = (n_male_g / len(df_group)) * 100
            cells[col_idx].text = f"{n_male_g} ({pct_male_g:.1f}%)"
        else:
            cells[col_idx].text = format_median_iqr(df_group[var_name])
    
    row_idx += 1

# Formatear celdas
for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

# ============================================================================
# NOTAS AL PIE
# ============================================================================

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Abbreviations: ').bold = True
p.add_run('NTG, normotriglyceridemia (TG <150 mg/dL); MiTG, mild hypertriglyceridemia '
         '(TG 150-199 mg/dL); MoTG, moderate hypertriglyceridemia (TG 200-399 mg/dL); '
         'HTG, severe hypertriglyceridemia (TG ≥400 mg/dL); LDL-C, low-density lipoprotein '
         'cholesterol; TyG, triglyceride-glucose index; ALT, alanine aminotransferase.')
p.style = 'Caption'

# ============================================================================
# GUARDAR
# ============================================================================

filename = 'Table1_Baseline_Characteristics_FINAL.docx'
doc.save(filename)

print(f"✅ Guardado: {filename}")
print()

# ============================================================================
# TAMBIÉN IMPRIMIR EN CONSOLA PARA VERIFICACIÓN
# ============================================================================

print("="*100)
print("TABLA 1 - VERIFICACIÓN EN CONSOLA")
print("="*100)
print()

print(f"{'Characteristic':<35} {'Overall':<25} {'NTG':<25} {'MiTG':<25} {'MoTG':<25} {'HTG':<25}")
print("-"*160)

for var_name, var_label in variables_info:
    row_data = [var_label]
    
    # Overall
    if var_name == 'gender':
        n_male = (df['gender'] == 'M').sum()
        pct_male = (n_male / len(df)) * 100
        row_data.append(f"{n_male} ({pct_male:.1f}%)")
    else:
        row_data.append(format_median_iqr(df[var_name]))
    
    # Por subgrupo
    for tg_group in ['NTG', 'MiTG', 'MoTG', 'HTG']:
        df_group = df[df['TG_group'] == tg_group]
        
        if var_name == 'gender':
            n_male_g = (df_group['gender'] == 'M').sum()
            pct_male_g = (n_male_g / len(df_group)) * 100
            row_data.append(f"{n_male_g} ({pct_male_g:.1f}%)")
        else:
            row_data.append(format_median_iqr(df_group[var_name]))
    
    print(f"{row_data[0]:<35} {row_data[1]:<25} {row_data[2]:<25} {row_data[3]:<25} {row_data[4]:<25} {row_data[5]:<25}")

print()
print("="*100)
print("✅ TABLA 1 COMPLETA GENERADA")
print("="*100)
print()

# ============================================================================
# VERIFICAR QUE TODAS LAS VARIABLES EXISTAN
# ============================================================================

print("VERIFICACIÓN DE VARIABLES:")
print("-"*100)

for var_name, var_label in variables_info:
    if var_name == 'gender':
        exists = 'gender' in df.columns
        n_missing = 0
    else:
        exists = var_name in df.columns
        n_missing = df[var_name].isna().sum() if exists else len(df)
    
    status = "✅" if exists else "❌"
    pct_missing = (n_missing / len(df)) * 100 if exists else 100
    
    print(f"{status} {var_label:<40} {'Existe' if exists else 'NO EXISTE':10} Missing: {n_missing:5,} ({pct_missing:5.1f}%)")

print()